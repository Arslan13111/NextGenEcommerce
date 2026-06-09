import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_instructions():
    doc = docx.Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('Instructions for Updating System Diagrams', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("This document provides exhaustive instructions for updating all architectural and design diagrams to align with the current NextGenEcommerce application. All diagrams must reflect the serverless architecture, Supabase integration, Tryona AI, and Snap Camera Kit integration.")

    # 1. Use Case Diagram
    doc.add_heading('1. Use Case Diagram Updates', level=1)
    doc.add_paragraph("Goal: Map all user roles and their interactions with the new cloud-based system.")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Element'
    hdr_cells[1].text = 'Instruction / Details'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Actors'
    row_cells[1].text = ("Primary Actors: Customer, Admin, Retailer, Delivery Personnel.\n"
                         "Secondary Actors (Systems): Supabase (Auth/DB), Tryona API, Snap Camera Kit, Payment Gateways (JazzCash, Safepay).")

    row_cells = table.add_row().cells
    row_cells[0].text = 'New Use Cases (Customer)'
    row_cells[1].text = ("- View Product in AR (Live Try-On)\n"
                         "- Generate AI Try-On (Static Image)\n"
                         "- Select JazzCash/EasyPaisa Payment\n"
                         "- Request Single-Item Return")

    row_cells = table.add_row().cells
    row_cells[0].text = 'Relationships'
    row_cells[1].text = ("- Use <<include>> for 'Login' in all checkout and profile actions.\n"
                         "- Use <<extend>> for 'AI Try-On' and 'AR Try-On' from 'View Product Details'.\n"
                         "- Connect 'Process Payment' to the JazzCash/Safepay secondary actors via an association line.")

    # 2. Entity Relationship Diagram (ERD)
    doc.add_heading('2. Entity Relationship Diagram (ERD) Updates', level=1)
    doc.add_paragraph("Goal: Align with the Supabase/PostgreSQL schema.")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    updates = [
        ("Product Table", "Add 'lens_id' (String), 'color_images' (JSONB/Map), 'ar_model_url' (String). Remove 'furniture_type'."),
        ("Order Table", "Add 'order_number' (Unique Text). Update 'payment_method' ENUM to include [JAZZCASH, EASYPAISA, SAFEPAY, COD]."),
        ("Order Items Table", "Ensure 1:N relationship with Orders. This is critical for the 'Single Item Return' logic."),
        ("User Table", "Add 'role' ENUM [customer, admin, retailer, delivery]. Add 'total_spent' (Decimal) to support the Tier system."),
        ("Relations", "Use Crow’s Foot notation. User (1) --- (N) Orders; Order (1) --- (N) Order_Items; Product (1) --- (N) Order_Items.")
    ]
    
    for item, instruction in updates:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = instruction

    # 3. Class Diagram
    doc.add_heading('3. Class Diagram (MVVM + Clean Architecture)', level=1)
    doc.add_paragraph("Goal: Reflect the Hilt-injected repositories and ViewModels.")
    
    classes = [
        ("ViewModels", "AuthViewModel, ProductViewModel, TryOnViewModel, OrderViewModel, CartViewModel. Each must have a 'StateFlow' for UI reactivity."),
        ("Repositories", "AuthRepository, ProductRepository, TryOnRepository (Tryona API), OrderRepository (Supabase), StorageRepository."),
        ("API Interfaces", "TryonaApiService (tryOnSimple), SafepayApiService, SupabaseClient."),
        ("Arrows", "Use standard UML arrows: Realization (dashed with hollow head) for Interface implementations; Composition (solid diamond) for ViewModels holding Repositories.")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for cls, desc in classes:
        row = table.add_row().cells
        row[0].text = cls
        row[1].text = desc

    # 4. Activity Diagram (Order & Return Flow)
    doc.add_heading('4. Activity Diagram: Order & Return', level=1)
    doc.add_paragraph("Must show the 'Order Splitting' logic.")
    doc.add_paragraph("- Start: Add items to Cart.\n"
                      "- Action: Checkout -> Select Payment Method.\n"
                      "- Decision: Is Payment Mobile (JazzCash)? If yes, 'Upload Receipt/Confirm' -> Set Status to 'Processing'.\n"
                      "- Fork: Split Cart Items into individual Orders.\n"
                      "- Join: Update Supabase Records.\n"
                      "- End: Order Success Screen.")

    # 5. Sequence Diagram (AI Try-On)
    doc.add_heading('5. Sequence Diagram: AI Try-On Process', level=1)
    doc.add_paragraph("Objects: User, TryOnScreen, TryOnViewModel, TryOnRepository, Tryona API.")
    doc.add_paragraph("1. User -> TryOnScreen: Select Image.\n"
                      "2. TryOnScreen -> TryOnViewModel: requestTryOn(personUri, garmentUrl).\n"
                      "3. TryOnViewModel -> TryOnRepository: processTryOn().\n"
                      "4. TryOnRepository -> Tryona API: POST /v1/tryon/simple (Multipart).\n"
                      "5. Tryona API -> TryOnRepository: Return image_url.\n"
                      "6. TryOnRepository -> TryOnViewModel: Result(Bitmap).\n"
                      "7. TryOnViewModel -> User: Display Result.")

    doc.save('C:\\Arslans Data\\FYP\\Diagram_Instructions\\Diagram_Update_Instructions.docx')

if __name__ == "__main__":
    create_instructions()
